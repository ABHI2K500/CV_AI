import os
import fitz 
import docx
import requests
from bs4 import BeautifulSoup
import groq
import docx
import PyPDF2
GROQ_API_KEY = "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"
groq_client = groq.Client(api_key="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")


def extract_text_from_cv(cv_path):
    """Extracts text from a given CV (PDF or DOCX)."""
    text = ""
    
    if cv_path.endswith(".docx"):
        doc = docx.Document(cv_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    
    elif cv_path.endswith(".pdf"):
        with open(cv_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

    return text.strip()


def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file."""
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file."""
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def get_job_description(job_url):
    """Attempts to scrape the job description from LinkedIn."""
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
    }\
    
    response = requests.get(job_url, headers=headers)
    if response.status_code != 200:
        return "Unable to fetch job description. Please paste it manually."

    soup = BeautifulSoup(response.text, "html.parser")
    
    
    job_desc = soup.find("div", class_="show-more-less-html__markup")
    
    return job_desc.get_text() if job_desc else "Job description not found. Please paste it manually."

import groq

groq_client = groq.Client(api_key="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")

def analyze_cv_vs_job(cv_text, job_desc):
    """Uses Groq AI to analyze the match between the CV and job description."""
    prompt = f"""
    Given the following CV and job description, analyze the match percentage.
    Provide:
    - A match score,(the match score indicates the percentage of getting the job, it is a must!!) (0-100%).
    - Key strengths of the CV.
    - Missing skills or qualifications.
    - Suggestions for improvement.

    CV:
    {cv_text}

    Job Description:
    {job_desc}

    Provide a detailed response.
    """

    response = groq_client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content

def main():
    cv_path = input("Enter the path of your CV (PDF or DOCX): ")
    
    if cv_path.endswith(".pdf"):
        cv_text = extract_text_from_pdf(cv_path)
    elif cv_path.endswith(".docx"):
        cv_text = extract_text_from_docx(cv_path)
    else:
        print("Unsupported file format.")
        return

    job_url = input("Enter the LinkedIn job URL: ")
    job_desc = get_job_description(job_url)

    if "Unable to fetch" in job_desc or "not found" in job_desc:
        job_desc = input("Please paste the job description manually: ")

    result = analyze_cv_vs_job(cv_text, job_desc)

    print("\n===== AI Analysis =====")
    print(result)

if __name__ == "__main__":
    main()
